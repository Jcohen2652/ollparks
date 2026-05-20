"""
Génération de documents à partir de templates Jinja2.

DOCX : python-docx (composition simple)
PDF  : weasyprint (rendu HTML/CSS)
HTML : direct
"""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

from api.services.documents.templates import TEMPLATES, TemplateSpec

TEMPLATES_DIR = Path(__file__).parent / "templates"


def _env() -> Environment:
    return Environment(
        loader=FileSystemLoader(str(TEMPLATES_DIR)),
        autoescape=select_autoescape(["html", "xml"]),
        trim_blocks=True,
        lstrip_blocks=True,
    )


def _resolve_required_fields(spec: TemplateSpec, data: dict[str, Any]) -> list[str]:
    """Retourne la liste des champs requis qui sont absents du payload."""
    missing: list[str] = []
    for field in spec.required_fields:
        cur: Any = data
        for part in field.split("."):
            if not isinstance(cur, dict) or part not in cur:
                missing.append(field)
                cur = None
                break
            cur = cur[part]
        if cur in (None, ""):
            missing.append(field)
    return missing


def render_template(template_key: str, data: dict[str, Any]) -> str:
    """Render HTML (sert pour PDF et HTML)."""
    spec = TEMPLATES.get(template_key)
    if not spec:
        raise ValueError(f"Template inconnu : {template_key}")
    missing = _resolve_required_fields(spec, data)
    if missing:
        raise ValueError(f"Champs requis manquants : {', '.join(missing)}")
    tpl = _env().get_template(f"{template_key}.html.j2")
    return tpl.render(**data, generated_at=datetime.utcnow().isoformat())


class DocumentGenerator:
    """High-level wrapper: render → bytes (HTML / PDF / DOCX)."""

    def render_html(self, template_key: str, data: dict[str, Any]) -> bytes:
        return render_template(template_key, data).encode("utf-8")

    def render_pdf(self, template_key: str, data: dict[str, Any]) -> bytes:
        # Import lazy : weasyprint a des deps lourdes (cairo, pango)
        from weasyprint import HTML

        html = render_template(template_key, data)
        buffer = io.BytesIO()
        HTML(string=html).write_pdf(buffer)
        return buffer.getvalue()

    def render_docx(self, template_key: str, data: dict[str, Any]) -> bytes:
        """Génère un DOCX simple (titre + paragraphes du HTML rendu)."""
        from html.parser import HTMLParser

        from docx import Document

        html = render_template(template_key, data)
        doc = Document()

        class TextExtractor(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self.lines: list[tuple[str, str]] = []  # (tag, text)
                self.current: list[str] = []
                self.tag = "p"

            def handle_starttag(self, tag: str, attrs: list) -> None:
                if tag in ("h1", "h2", "h3", "p", "li"):
                    self._flush()
                    self.tag = tag

            def handle_endtag(self, tag: str) -> None:
                if tag in ("h1", "h2", "h3", "p", "li"):
                    self._flush()
                    self.tag = "p"

            def handle_data(self, data: str) -> None:
                if data.strip():
                    self.current.append(data)

            def _flush(self) -> None:
                if self.current:
                    self.lines.append((self.tag, " ".join(self.current).strip()))
                    self.current = []

            def close(self) -> None:
                self._flush()
                super().close()

        ext = TextExtractor()
        ext.feed(html)
        ext.close()

        for tag, text in ext.lines:
            if tag == "h1":
                doc.add_heading(text, level=1)
            elif tag == "h2":
                doc.add_heading(text, level=2)
            elif tag == "h3":
                doc.add_heading(text, level=3)
            else:
                doc.add_paragraph(text)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
